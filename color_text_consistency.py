#!/usr/bin/python2
# -*- coding: utf-8 -*-
"""
Color code for Word documents (DOCX): Colors of consciousness
O.Colizoli 2025
Python 3.9

Notes
-------------
The following packages need to be installed: python-docx, pandas, numpy
>> pip install python-docx pandas numpy
---
To run:
>> python color_text_consistency.py
You will be prompted with the book name and whether you want to change the font type and size.
---
This script will loop through individual letters and reformat each occurrence of a letter with a specific color (RGB);
Each letter is also assigned a consistency value: e.g. if "a" is 100% consistent, then all "a"s will receive the same color. 
If "a" is 75% consistent, then it will receive a different color with 25% probability. 
The other colors are chosen from a uniform distribution.

IMPORTANT!!
The order of the letters in the probability_distributions.csv files 
need to match the order of the letters in the sub-xxx_prediction_space.csv files!

If 'change_color' is True, the script will run the coloring-consistency process described above. 
If 'change_font' is True, the script will replace the entire document's font typeface and font size. 
The reformatted document is saved as a new document. 
Fonts need to be already installed on the computer.
---
Books are found in the directory: os.path.join(os.getcwd(), 'books')
The name of book input at the prompt should exclude the 'docx' file extension: "books/sub-{}_book{}.docx".format(subject, book_number)
The new colored version will be saved as: "books/{}_processed.docx".format(book_name)
---
Letters and colors to be changed are imported from a CSV file in the directory: os.path.join(os.getcwd(), 'colors')
The CSV file with colors needs to have the following columns: 'letter', 'rgb_r', 'rgb_g', 'rbg_b'
Letters are case-sensitive.
---
Character formatting is applied at the docx.text.run.Run level. 
The script can be adjusted to change the font typeface, size, bold, italic, 
and underline of single letters or the whole document.
A Run object has a read-only font property providing access to a Font object. 
A run's Font object (docx.text.run.Run.font) provides properties for getting and setting the character formatting for that run.
E.g. current_run.font.color.rgb = RGBColor(rgb_r, rgb_g, rbg_b) 
---
The function for isolating individual letters as runs, isolate_run(), was taken from here:
See: https://github.com/python-openxml/python-docx/issues/980
"""

import os, time, itertools, copy
import docx
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import pandas as pd
import numpy as np

from IPython import embed as shell # for Olympia's debugging only


def isolate_run(paragraph, start, end):
    """Return docx.text.run.Run object containing only `paragraph.text[start:end]`.
    
    Notes
    -----
    Runs are split as required to produce a new run at the `start` that ends at `end`.
    Runs are unchanged if the indicated range of text already occupies its own run. The
    resulting run object is returned.

    `start` and `end` are as in Python slice notation. For example, the first three
    characters of the paragraph have (start, end) of (0, 3). `end` is not the index of
    the last character. These correspond to `match.start()` and `match.end()` of a regex
    match object and `s[start:end]` of Python slice notation.
    
    https://github.com/python-openxml/python-docx/issues/980
    """
    rs = tuple(paragraph._p.r_lst)

    def advance_to_run_containing_start(start, end):
        """Return (r_idx, start, end) triple indicating start run and adjusted offsets.

        The start run is the run the `start` offset occurs in. The returned `start` and
        `end` values are adjusted to be relative to the start of `r_idx`.
        """
        # --- add 0 at end so `r_ends[-1] == 0` ---
        r_ends = tuple(itertools.accumulate(len(r.text) for r in rs)) + (0,)
        r_idx = 0
        while start >= r_ends[r_idx]:
            r_idx += 1
        skipped_rs_offset = r_ends[r_idx - 1]
        return rs[r_idx], r_idx, start - skipped_rs_offset, end - skipped_rs_offset

    def split_off_prefix(r, start, end):
        """Return adjusted `end` after splitting prefix off into separate run.

        Does nothing if `r` is already the start of the isolated run.
        """
        if start > 0:
            prefix_r = copy.deepcopy(r)
            r.addprevious(prefix_r)
            r.text = r.text[start:]
            prefix_r.text = prefix_r.text[:start]
        return end - start

    def split_off_suffix(r, end):
        """Split `r` at `end` such that suffix is in separate following run."""
        suffix_r = copy.deepcopy(r)
        r.addnext(suffix_r)
        r.text = r.text[:end]
        suffix_r.text = suffix_r.text[end:]

    def lengthen_run(r, r_idx, end):
        """Add prefixes of following runs to `r` until `end` is reached."""
        while len(r.text) < end:
            suffix_len_reqd = end - len(r.text)
            r_idx += 1
            next_r = rs[r_idx]
            if len(next_r.text) <= suffix_len_reqd:
                # --- subsume next run ---
                r.text = r.text + next_r.text
                next_r.getparent().remove(next_r)
                continue
            if len(next_r.text) > suffix_len_reqd:
                # --- take prefix from next run ---
                r.text = r.text + next_r.text[:suffix_len_reqd]
                next_r.text = next_r.text[suffix_len_reqd:]

    r, r_idx, start, end = advance_to_run_containing_start(start, end)
    end = split_off_prefix(r, start, end)

    # --- if run is longer than isolation-range we need to split-off a suffix run ---
    if len(r.text) > end:
        split_off_suffix(r, end)
    # --- if run is shorter than isolation-range we need to lengthen it by taking text
    # --- from subsequent runs
    elif len(r.text) < end:
        lengthen_run(r, r_idx, end)


    return docx.text.run.Run(r, paragraph)


def replace_letters_with_colors():
    '''In a Microsoft Word document (DOCX), reformat each occurrence of a letter with a specific color (RGB).
    
    Notes:
    ------
    See notes at top of script for more information. 
    '''
    # subject = input('Subject Number: ')
    # book_number = input('Book Number: ')
    # change_color = int(input('Change colors? (1 for Yes, 0 for No): '))
    # change_font = int(input('Change Font? (1 for Yes, 0 for No): '))
    subject = '01'
    book_number = '1'
    change_color = True
    change_font = False
    
    if change_font:
        # replace_font = input('Font name: ')
        # replace_size = np.float32(input('Font size: '))
        replace_font = 'Arial Black'
        replace_size = 11

    in_book_filename = os.path.join('books', 'sub-{}_book{}.docx'.format(subject, book_number)) # original
    out_book_filename = os.path.join('books', 'sub-{}_book{}_processed.docx'.format(subject, book_number)) # save as new
    
    if change_color:
        # Define letters and colors to replace
        df_letters = pd.read_csv(os.path.join('colors', 'sub-{}_letter_colour_pairs_sorted.csv'.format(subject))) # the letters and their consistency conditions
        df_colors = pd.read_csv(os.path.join('colors', 'rgb_colors.csv')) # the CSV file with the 'letters' and 'r', 'g', 'b' values
        # join these two dataframes on the color numbers        
        df = df_letters.merge(df_colors, how='inner', on='colour_id')
        
        # get only train==yes letters
        df = df[df['train']=="yes"].copy()
        df.reset_index(inplace=True)

        letters = df['letter'] # vowels including y
        
        # determine which letter set this participant has for the probability distributions
        if "e" in np.array(letters):
            letter_set = "set1"
            order =[ 'e', 's', 'm', 'q', 'x', 'c', 'h', 'o'] # needs to match the order of letters in the probabilit_distributions file
        elif "a" in np.array(letters):
            letter_set = "set2"
            order = ['a', 'n', 'w', 'z', 'j', 'f', 'r', 'i'] # needs to match the order of letters in the probabilit_distributions file

        else:
            print("ERROR: check letter sets! Letters not in either set.")
        
        # make 'letters' a categorical column with that order
        df['letter'] = pd.Categorical(df['letter'], categories=order, ordered=True)
        # sort by 'letters' using that order
        df = df.sort_values('letter').reset_index(drop=True)
        
        ### DOUBLE CHECK THIS!! ###
        # rows are letters, columns are the corresponding letter's colorcode index
        prob_dist = pd.read_csv(os.path.join('colors', 'probability_distributions_{}.csv'.format(letter_set))) 
        color_array = df['colour_id'] # array of color codes in order of the letters for this participant
        
        # First, run the letter-by-letter search and replace loop
        doc = Document(in_book_filename)
        for idx_letter,letter in enumerate(letters): # loop over letters
            print(letter)
            # get corresponding probability distribution
            p = prob_dist[letter]
            
            print('Searching for "{}"...'.format(letter))
            print('Number of paragraphs: {}'.format(len(doc.paragraphs)))
            for p_idx,paragraph in enumerate(doc.paragraphs):
                # print('Paragraph {}'.format(p_idx))
                for start in range(len(paragraph.text)): # isolate runs that are 1 character in length only
                    # Here the "1" indicates the step size to search for strings, i.e., ending only 1 unit from start of string will return a single character in a single run
                    end = start + 1 
                    current_run = isolate_run(paragraph, start, end)
                    if current_run.text == letter: # only change the color of letters in the CSV file (case-sensitive)
                        # draw color based on distribution: numpy.random.choice(a, p)
                        colorcode = np.random.choice(a = np.array(color_array), p = np.array(p))
                        # current_run.font.color.rgb = RGBColor(int(df_colors['rgb_r'][colorcode-1]), int(df_colors['rgb_g'][colorcode-1]), int(df_colors['rgb_b'][colorcode-1]))
                        current_run.font.color.rgb = RGBColor(
                            int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_r'].iloc[0]), 
                            int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_g'].iloc[0]), 
                            int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_b'].iloc[0]), 
                            )   
                                         
                        print(colorcode)
                        # print(
                        #     int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_r'].iloc[0]),
                        #     int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_g'].iloc[0]),
                        #     int(df_colors.loc[df_colors['colour_id'] == colorcode, 'rgb_b'].iloc[0]),
                        # )
                        # If you want to change other formatting options of the individual letters, you can specify that here:
                        # current_run.font.size
                        # current_run.font.name 
                        # current_run.font.italic
                        # etc...


    # Replace the entire doc's font typeface (name) and size
    if change_font: 
        print('Changing font to {} and size {}...'.format(replace_font, replace_size))
        for paragraph in doc.paragraphs:
            for r in paragraph.runs:
                r.font.name = replace_font
                r.font.size = Pt(replace_size)
                #r.font.color.rgb = RGBColor(0, 0, 0) # e.g., set all to black
    # save as new book
    doc.save(out_book_filename)
    print('New book saved as {}'.format(out_book_filename))
    

''' RUN '''      
if __name__ == '__main__':
    t0 = time.time() # measure run time (not optimized, just curious)
    replace_letters_with_colors()
    print('It took {} minutes'.format( (time.time()-t0)/60 )) # report run time
    
    
    

